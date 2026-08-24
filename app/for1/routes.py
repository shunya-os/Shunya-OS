"""FOR-1 Routes — Proposal Intelligence Engine API + UI."""

import json
import uuid
import os
from datetime import datetime, timezone
from flask import jsonify, render_template, request, session, url_for, current_app, redirect

from app import db
from app.for1 import for1_bp
from app.for1.engine import (
    generate_proposal,
    search_knowledge,
    search_suppliers,
    render_proposal_html,
    generate_proposal_pdf,
)
from app.models import Proposal, ProposalVersion, KnowledgeDocument, Lead
from app.tenant import Tenant
from app.relationship.integration import record_event, update_ai_memory_from_event


# ── Helper ──────────────────────────────────────────────────────────────

def _get_tenant():
    """Get the current tenant from session or return None."""
    from app.tenant import Tenant
    tenant_id = session.get("tenant_id")
    if tenant_id:
        return db.session.get(Tenant, tenant_id)
    return None


def _get_organization():
    """Get the canonical Organization from session."""
    from app.models import Organization
    org_id = session.get("current_org_id")
    if org_id:
        return db.session.get(Organization, org_id)
    return None


def _require_auth():
    """Check authentication. Returns redirect for HTML, None if OK."""
    if not session.get("user_id"):
        return redirect(url_for("auth.login_page", next=request.path))
    return None


def _api_require_auth():
    """Check authentication for API routes. Returns JSON 401 or None."""
    if not session.get("user_id"):
        return jsonify({"error": "Authentication required"}), 401
    return None


# ── HTML Pages ──────────────────────────────────────────────────────────


@for1_bp.route("/for1/proposals")
def proposal_list():
    """List all proposals."""
    auth_err = _require_auth()
    if auth_err:
        return auth_err
    proposals = Proposal.query.order_by(Proposal.created_at.desc()).all()
    return render_template("for1_proposals.html", proposals=proposals)


@for1_bp.route("/for1/proposal/<int:proposal_id>")
def proposal_view(proposal_id: int):
    """View a single proposal."""
    auth_err = _require_auth()
    if auth_err:
        return auth_err
    proposal = db.session.get(Proposal, proposal_id)
    if not proposal:
        return "Proposal not found", 404
    tenant = _get_tenant()
    return render_template("for1_proposal.html", proposal=proposal, tenant=tenant)


@for1_bp.route("/for1/proposal/<int:proposal_id>/preview")
def proposal_preview(proposal_id: int):
    """Preview the proposal as a beautiful web page."""
    proposal = db.session.get(Proposal, proposal_id)
    if not proposal:
        return "Proposal not found", 404
    tenant = _get_tenant()
    html = render_proposal_html(proposal, tenant)
    return html


@for1_bp.route("/for1/dashboard")
def for1_dashboard():
    """FOR-1 Executive dashboard."""
    auth_err = _require_auth()
    if auth_err:
        return auth_err
    now = datetime.now(timezone.utc)
    # Quick stats
    total_leads = Lead.query.count()
    pending_proposals = Proposal.query.filter_by(status="draft").count()
    today_leads = Lead.query.filter(
        Lead.created_at >= now.replace(hour=0, minute=0, second=0)
    ).count()
    recent_leads = Lead.query.order_by(Lead.created_at.desc()).limit(10).all()
    recent_proposals = Proposal.query.order_by(Proposal.created_at.desc()).limit(10).all()
    return render_template(
        "for1_dashboard.html",
        total_leads=total_leads,
        pending_proposals=pending_proposals,
        today_leads=today_leads,
        recent_leads=recent_leads,
        recent_proposals=recent_proposals,
    )


# ── API Routes ──────────────────────────────────────────────────────────


@for1_bp.route("/api/v1/for1/proposals", methods=["GET"])
def api_proposal_list():
    """List all proposals (JSON API)."""
    auth_err = _require_auth()
    if auth_err:
        return auth_err
    proposals = Proposal.query.order_by(Proposal.created_at.desc()).all()
    return jsonify({"proposals": [p.to_dict() if hasattr(p, 'to_dict') else {
        "id": p.id, "title": p.title or f"Proposal #{p.id}", "status": p.status or "draft",
        "budget": float(p.budget) if p.budget else 0, "created_at": p.created_at.isoformat() if p.created_at else None
    } for p in proposals]})


@for1_bp.route("/api/v1/for1/proposals", methods=["POST"])
def api_create_proposal():
    """Create a new proposal from lead data, optionally with AI generation."""
    auth_err = _api_require_auth()
    if auth_err:
        return auth_err

    data = request.get_json(silent=True) or {}
    lead_id = data.get("lead_id")
    if not lead_id:
        return jsonify({"error": "lead_id is required"}), 400

    lead = db.session.get(Lead, lead_id)
    if not lead:
        return jsonify({"error": "Lead not found"}), 404

    tenant = _get_tenant()
    organization = _get_organization()
    org_id = organization.id if organization else (getattr(tenant, "id", None) or None)
    proposal = Proposal(
        organization_id=org_id,
        opportunity_id=lead_id,
        relationship_id=data.get("relationship_id"),
        title=data.get("title", f"Proposal for {lead.customer_name or 'Client'}"),
        destination=lead.destination or data.get("destination", ""),
        pax=lead.pax or data.get("pax", ""),
        budget=lead.budget or data.get("budget", 0),
        status="draft",
        created_by=session.get("user_id", ""),
    )

    # If AI generation requested
    if data.get("ai_generate"):
        proposal.status = "ai_generating"
        db.session.add(proposal)
        db.session.commit()

        lead_data = {
            "customer_name": lead.customer_name,
            "destination": lead.destination,
            "pax": lead.pax,
            "dates": lead.dates,
            "budget": float(lead.budget or 0),
            "notes": lead.notes,
        }
        knowledge_docs = search_knowledge(f"{lead.destination} travel", getattr(tenant, "id", 0))
        suppliers = search_suppliers(lead.destination or "", getattr(tenant, "id", 0))
        result = generate_proposal(lead_data, tenant, knowledge_docs, suppliers)

        if "error" in result:
            proposal.status = "draft"
            proposal.generation_notes = result["error"]
        else:
            proposal.title = result.get("title", proposal.title)
            proposal.destination = result.get("destination", proposal.destination)
            proposal.duration_days = result.get("duration_days", 0)
            proposal.itinerary_json = json.dumps(result.get("itinerary", []))
            proposal.pricing_json = json.dumps(result.get("pricing", {}))
            proposal.inclusions = "\n".join(result.get("inclusions", []))
            proposal.exclusions = "\n".join(result.get("exclusions", []))
            proposal.terms = result.get("terms", "")
            proposal.ai_generated = True
            proposal.ai_model = "openai/gpt-4o-mini"
            proposal.ai_prompt = json.dumps(lead_data)
            proposal.status = "review"

        # Generate web HTML
        proposal.web_html = render_proposal_html(proposal, tenant)

        db.session.commit()

        # Create version 1 snapshot
        version = ProposalVersion(
            proposal_id=proposal.id,
            version_number=1,
            snapshot_json=json.dumps(proposal.to_dict()),
            change_summary="AI-generated initial proposal",
            created_by=session.get("user_id", ""),
        )
        db.session.add(version)
        db.session.commit()

        # Record timeline event on associated Relationship
        if proposal.relationship_id and proposal.organization_id:
            record_event(
                relationship_id=proposal.relationship_id,
                organization_id=proposal.organization_id,
                event_type="proposal.generated",
                title=f"Proposal generated: {proposal.title}",
                description=f"Status: {proposal.status}, Amount: {proposal.currency} {proposal.budget}",
                reference_type="proposal",
                reference_id=proposal.id,
                created_by=session.get("user_id", ""),
            )
            update_ai_memory_from_event(
                relationship_id=proposal.relationship_id,
                organization_id=proposal.organization_id,
                event_type="proposal.generated",
                summary_fragment=f"Proposal generated for {proposal.destination}: {proposal.title}",
            )
            db.session.commit()

        return jsonify({"success": True, "proposal": proposal.to_dict()}), 201

    db.session.add(proposal)
    db.session.commit()
    return jsonify({"success": True, "proposal": proposal.to_dict()}), 201


@for1_bp.route("/api/v1/for1/proposals/<int:proposal_id>", methods=["GET"])
def api_get_proposal(proposal_id: int):
    """Get a proposal by ID."""
    proposal = db.session.get(Proposal, proposal_id)
    if not proposal:
        return jsonify({"error": "Proposal not found"}), 404
    return jsonify({"success": True, "proposal": proposal.to_dict(include_html=True)})


@for1_bp.route("/api/v1/for1/proposals/<int:proposal_id>/pdf", methods=["POST"])
def api_generate_pdf(proposal_id: int):
    """Generate PDF for a proposal."""
    proposal = db.session.get(Proposal, proposal_id)
    if not proposal:
        return jsonify({"error": "Proposal not found"}), 404
    tenant = _get_tenant()
    pdf_path = generate_proposal_pdf(proposal, tenant)
    if pdf_path.startswith("<!--"):
        return jsonify({"error": pdf_path}), 500
    proposal.pdf_path = pdf_path
    db.session.commit()
    return jsonify({"success": True, "pdf_url": pdf_path})


@for1_bp.route("/api/v1/for1/proposals/<int:proposal_id>/status", methods=["PATCH"])
def api_update_proposal_status(proposal_id: int):
    """Update proposal status (send, accept, cancel)."""
    proposal = db.session.get(Proposal, proposal_id)
    if not proposal:
        return jsonify({"error": "Proposal not found"}), 404
    data = request.get_json(silent=True) or {}
    new_status = data.get("status")
    if new_status not in ("review", "sent", "accepted", "booked", "cancelled"):
        return jsonify({"error": f"Invalid status: {new_status}"}), 400

    old_status = proposal.status
    proposal.status = new_status
    if new_status == "sent":
        proposal.sent_at = datetime.now(timezone.utc)
        proposal.sent_via = data.get("sent_via", "link")
    elif new_status == "accepted":
        proposal.accepted_at = datetime.now(timezone.utc)

    db.session.commit()

    # Create version snapshot
    version = ProposalVersion(
        proposal_id=proposal.id,
        version_number=proposal.version_number + 1,
        snapshot_json=json.dumps(proposal.to_dict()),
        change_summary=f"Status changed: {old_status} → {new_status}",
        created_by=session.get("user_id", ""),
    )
    db.session.add(version)
    proposal.version_number += 1
    db.session.commit()

    return jsonify({"success": True, "proposal": proposal.to_dict()})


@for1_bp.route("/api/v1/for1/proposals/<int:proposal_id>/regenerate", methods=["POST"])
def api_regenerate_proposal(proposal_id: int):
    """Regenerate a proposal with AI (creates new version)."""
    auth_err = _api_require_auth()
    if auth_err:
        return auth_err
    proposal = db.session.get(Proposal, proposal_id)
    if not proposal:
        return jsonify({"error": "Proposal not found"}), 404
    lead = db.session.get(Lead, proposal.lead_id)
    if not lead:
        return jsonify({"error": "Lead not found"}), 404

    tenant = _get_tenant()
    data = request.get_json(silent=True) or {}
    lead_data = {
        "customer_name": lead.customer_name,
        "destination": proposal.destination or lead.destination,
        "pax": proposal.pax or lead.pax,
        "dates": lead.dates,
        "budget": float(data.get("budget", proposal.budget or lead.budget or 0)),
        "notes": data.get("notes", lead.notes),
    }

    proposal.status = "ai_generating"
    db.session.commit()

    knowledge_docs = search_knowledge(f"{proposal.destination} travel", getattr(tenant, "id", 0))
    suppliers = search_suppliers(proposal.destination or "", getattr(tenant, "id", 0))
    result = generate_proposal(lead_data, tenant, knowledge_docs, suppliers)

    if "error" in result:
        proposal.status = "draft"
        proposal.generation_notes = result["error"]
        db.session.commit()
        return jsonify({"error": result["error"]}), 500

    proposal.title = result.get("title", proposal.title)
    proposal.duration_days = result.get("duration_days", proposal.duration_days)
    proposal.itinerary_json = json.dumps(result.get("itinerary", []))
    proposal.pricing_json = json.dumps(result.get("pricing", {}))
    proposal.inclusions = "\n".join(result.get("inclusions", []))
    proposal.exclusions = "\n".join(result.get("exclusions", []))
    proposal.terms = result.get("terms", "")
    proposal.ai_generated = True
    proposal.status = "review"
    proposal.web_html = render_proposal_html(proposal, tenant)
    proposal.version_number += 1

    db.session.commit()

    version = ProposalVersion(
        proposal_id=proposal.id,
        version_number=proposal.version_number,
        snapshot_json=json.dumps(proposal.to_dict()),
        change_summary="AI regenerated proposal",
        created_by=session.get("user_id", ""),
    )
    db.session.add(version)
    db.session.commit()

    return jsonify({"success": True, "proposal": proposal.to_dict()}), 200


# ── Knowledge Document API ──────────────────────────────────────────────


ALLOWED_EXTENSIONS = {"pdf", "docx", "xlsx", "png", "jpg", "jpeg", "mp4", "mov"}


def _extract_text_from_file(filepath: str, file_type: str) -> str:
    """Extract text content from an uploaded file."""
    text = ""
    try:
        if file_type == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif file_type == "docx":
            import docx
            doc = docx.Document(filepath)
            text = "\n".join(p.text for p in doc.paragraphs)
        elif file_type in ("xlsx",):
            import openpyxl
            wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    text += " ".join(str(c) for c in row if c) + "\n"
        elif file_type in ("png", "jpg", "jpeg"):
            from PIL import Image
            img = Image.open(filepath)
            text = f"[Image: {os.path.basename(filepath)}, {img.size[0]}x{img.size[1]}, {img.mode}]"
    except Exception as e:
        text = f"[Extraction error: {e}]"
    return text[:100000]  # cap at 100K chars


@for1_bp.route("/api/v1/for1/knowledge/upload", methods=["POST"])
def api_upload_knowledge():
    """Upload a knowledge document."""
    auth_err = _api_require_auth()
    if auth_err:
        return auth_err

    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "Empty filename"}), 400

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({"error": f"File type .{ext} not supported. Supported: {', '.join(ALLOWED_EXTENSIONS)}"}), 400

    # Save file
    upload_dir = os.path.join(current_app.root_path, "..", "static", "uploads", "knowledge")
    os.makedirs(upload_dir, exist_ok=True)
    safe_name = f"{uuid.uuid4().hex[:12]}_{file.filename}"
    filepath = os.path.join(upload_dir, safe_name)
    file.save(filepath)

    title = request.form.get("title", file.filename)
    category = request.form.get("category", "general")
    tags = request.form.get("tags", "")

    extracted_text = _extract_text_from_file(filepath, ext)
    summary = extracted_text[:500] if extracted_text else ""

    tenant = _get_tenant()
    doc = KnowledgeDocument(
        tenant_id=getattr(tenant, "id", None),
        title=title,
        category=category,
        file_path=f"/static/uploads/knowledge/{safe_name}",
        file_size_bytes=os.path.getsize(filepath),
        extracted_text=extracted_text,
        summary=summary,
        tags=tags,
        uploaded_by=session.get("user_id", ""),
    )
    db.session.add(doc)
    db.session.commit()

    return jsonify({"success": True, "document": doc.to_dict()}), 201


@for1_bp.route("/api/v1/for1/knowledge", methods=["GET"])
def api_list_knowledge():
    """List knowledge documents with optional search."""
    auth_err = _api_require_auth()
    if auth_err:
        return auth_err

    query = request.args.get("q", "").strip()
    if query:
        results = search_knowledge(query)
        return jsonify({"success": True, "documents": results})

    documents = KnowledgeDocument.query.order_by(KnowledgeDocument.created_at.desc()).all()
    return jsonify({"success": True, "documents": [d.to_dict() for d in documents]})


@for1_bp.route("/api/v1/for1/knowledge/<int:doc_id>", methods=["DELETE"])
def api_delete_knowledge(doc_id: int):
    """Delete a knowledge document."""
    auth_err = _api_require_auth()
    if auth_err:
        return auth_err
    doc = db.session.get(KnowledgeDocument, doc_id)
    if not doc:
        return jsonify({"error": "Document not found"}), 404
    db.session.delete(doc)
    db.session.commit()
    return jsonify({"success": True})


@for1_bp.route("/api/v1/for1/leads", methods=["GET"])
def api_lead_list():
    """List all leads (JSON API)."""
    auth_err = _api_require_auth()
    if auth_err:
        return auth_err
    leads = Lead.query.order_by(Lead.created_at.desc()).all()
    return jsonify({"leads": [{
        "id": l.id, "customer_name": l.customer_name or f"Lead #{l.id}",
        "status": l.status or "new", "created_at": l.created_at.isoformat() if l.created_at else None
    } for l in leads]})


@for1_bp.route("/api/v1/for1/leads", methods=["POST"])
def api_lead_create():
    """Create a new lead (JSON API)."""
    auth_err = _api_require_auth()
    if auth_err:
        return auth_err
    data = request.get_json(silent=True) or {}
    if not data.get("customer_name"):
        return jsonify({"error": "customer_name required"}), 400
    lead = Lead(customer_name=data["customer_name"], status="new")
    lead.code = f"L{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}{uuid.uuid4().hex[:4]}"
    lead.source = data.get("source", "api")
    db.session.add(lead)
    db.session.commit()
    return jsonify({"id": lead.id, "success": True}), 201


# ── Lead → Proposal quick-create ────────────────────────────────────────


@for1_bp.route("/api/v1/for1/leads/<int:lead_id>/proposal", methods=["POST"])
def api_lead_to_proposal(lead_id: int):
    """Quick-create a proposal from a lead with optional AI."""
    auth_err = _api_require_auth()
    if auth_err:
        return auth_err
    data = request.get_json(silent=True) or {}
    data["lead_id"] = lead_id
    # Proxy to the create proposal endpoint
    return api_create_proposal()