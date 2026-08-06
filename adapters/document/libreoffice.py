"""
LibreOffice Document Adapter.

Creates, edits, converts, exports, and imports documents.
Backed by ``libreoffice --headless`` when available; falls back to
python-docx (``.docx``) and fpdf2 (``.pdf``) for a fully functional
implementation without any external dependency.
"""

from __future__ import annotations

import io
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Any

from adapters import DocumentAdapter


# ---------------------------------------------------------------------------
# Detection
# ---------------------------------------------------------------------------

def _libreoffice_available() -> bool:
    """Return True if ``libreoffice`` (or ``soffice``) is on PATH."""
    return (
        shutil.which("libreoffice") is not None
        or shutil.which("soffice") is not None
    )


_HAVE_LO = _libreoffice_available()
_HAVE_DOCX = True  # python-docx is bundled
_HAVE_FPDF = True  # fpdf2 is bundled


# ---------------------------------------------------------------------------
# Adapter
# ---------------------------------------------------------------------------

class LibreOfficeAdapter(DocumentAdapter):
    """Document creation, editing, conversion, and text extraction.

    Uses ``libreoffice --headless`` when installed (recommended).
    Falls back to pure-Python libraries (python-docx, fpdf2, zip-based ODF)
    when the system CLI is unavailable.
    """

    # ── lifecycle ──────────────────────────────────────────────────────

    def __init__(self) -> None:
        self._lo_available = _HAVE_LO

    # ── Public API ─────────────────────────────────────────────────────

    def create_document(
        self,
        title: str,
        content: str,
        fmt: str = "odt",
    ) -> str:
        """Create a new document and return its file path.

        Supported output formats:
            ``odt``  — ODF text document (LibreOffice native)
            ``docx`` — Office Open XML (Word)
            ``pdf``  — Portable Document Format
            ``txt``  — Plain text
            ``html`` — HTML fragment
        """
        fmt = fmt.lower()
        if fmt not in _SUPPORTED_CREATE:
            raise ValueError(
                f"LibreOfficeAdapter.create_document supports "
                f"{sorted(_SUPPORTED_CREATE)}, got {fmt!r}"
            )

        out_dir = tempfile.mkdtemp(prefix="shunya_lo_")
        out_path = os.path.join(out_dir, f"{_safe_filename(title)}.{fmt}")

        creators: dict[str, Any] = {
            "odt": _create_odt,
            "docx": _create_docx,
            "pdf": _create_pdf,
            "txt": _create_txt,
            "html": _create_html,
        }
        creators[fmt](out_path, title, content)
        return out_path

    def edit_document(
        self,
        path: str,
        new_content: str | None = None,
        insert_at: str | int | None = None,
        **kwargs: Any,
    ) -> str:
        """Edit an existing document.

        Parameters
        ----------
        path : str
            Path to the document to edit.
        new_content : str, optional
            If given, replaces the entire document body content.
        insert_at : str or int, optional
            If ``"end"`` or an integer paragraph index, the *new_content* is
            appended at that position rather than replacing everything.
        **kwargs
            Additional keyword arguments (reserved for future use).

        Returns
        -------
        str
            Path to the edited document (same file, overwritten in place).
        """
        src = Path(path)
        if not src.is_file():
            raise FileNotFoundError(f"Document not found: {path}")

        ext = src.suffix.lower()
        content = new_content or ""

        if ext == ".docx":
            _edit_docx(src, content, insert_at)
        elif ext == ".odt":
            _edit_odt(src, content, insert_at)
        elif ext == ".txt":
            _edit_txt(src, content, insert_at)
        else:
            # For PDF and others: extract, rebuild as docx, convert back
            text = self.extract_text(path)
            if new_content is not None:
                text = new_content
            elif insert_at is not None:
                text = text + "\n" + content
            tmp = self.create_document("edited", text, fmt="docx")
            converted = self.convert(tmp, ext.lstrip("."))
            shutil.copy2(converted, str(src))
        return str(src.resolve())

    def convert(self, source_path: str, target_fmt: str) -> str:
        """Convert a document to *target_fmt*.

        Supports round-trip between ODT, DOCX, PDF, TXT, and HTML.
        Uses LibreOffice CLI when available, otherwise a pure-Python
        DOCX→PDF pipeline.
        """
        target_fmt = target_fmt.lower()
        src = Path(source_path)
        if not src.is_file():
            raise FileNotFoundError(f"Source file not found: {source_path}")

        out_dir = tempfile.mkdtemp(prefix="shunya_lo_convert_")

        # ── LibreOffice path ──────────────────────────────────────
        if self._lo_available and target_fmt in _LO_CONVERT_MAP:
            filter_name = _LO_CONVERT_MAP[target_fmt]
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to", f"{target_fmt}:{filter_name}",
                "--outdir", out_dir,
                str(src.resolve()),
            ]
            try:
                subprocess.run(cmd, capture_output=True, timeout=120, check=True)
            except FileNotFoundError:
                self._lo_available = False  # downgrade gracefully
                return self.convert(source_path, target_fmt)
            except subprocess.CalledProcessError as exc:
                msg = exc.stderr.decode() if exc.stderr else str(exc)
                raise RuntimeError(f"LibreOffice conversion failed: {msg}") from exc

            result = os.path.join(out_dir, f"{src.stem}.{target_fmt}")
            if os.path.isfile(result):
                return result
            candidates = list(Path(out_dir).iterdir())
            if candidates:
                return str(candidates[0])
            raise RuntimeError(
                f"LibreOffice conversion produced no output in {out_dir}"
            )

        # ── Pure-Python fallback ──────────────────────────────────
        return _convert_pure(src, target_fmt, out_dir)

    def extract_text(self, path: str) -> str:
        """Extract plain text from a document.

        Supports ODT, DOCX, PDF, TXT, and HTML files.
        """
        src = Path(path)
        if not src.is_file():
            raise FileNotFoundError(f"File not found: {path}")

        ext = src.suffix.lower()

        if ext == ".docx":
            return _extract_docx(src)
        elif ext == ".odt":
            return _extract_odt(src)
        elif ext == ".pdf":
            return _extract_pdf(src)
        elif ext == ".txt":
            return src.read_text(encoding="utf-8", errors="replace")
        elif ext in (".html", ".htm"):
            return _extract_html(src)
        else:
            # Try LibreOffice TXT conversion as a last resort
            if self._lo_available:
                txt_path = self.convert(path, "txt")
                return Path(txt_path).read_text(encoding="utf-8", errors="replace")
            raise ValueError(f"Unsupported format for text extraction: {ext}")


# ===================================================================
# Supported-format registries
# ===================================================================

_SUPPORTED_CREATE: set[str] = {"odt", "docx", "pdf", "txt", "html"}

_LO_CONVERT_MAP: dict[str, str] = {
    "odt": "writer8",
    "docx": "MS Word 2007 XML",
    "pdf": "writer_pdf_Export",
    "txt": "Text",
    "html": "HTML",
    "rtf": "Rich Text Format",
}


# ===================================================================
# Pure-Python conversion fallback
# ===================================================================

def _convert_pure(src: Path, target_fmt: str, out_dir: str) -> str:
    """Convert *src* to *target_fmt* using only pure-Python libraries."""
    ext = src.suffix.lower()
    stem = src.stem
    out_path = os.path.join(out_dir, f"{stem}.{target_fmt}")

    # ── Same-format passthrough ────────────────────────────────────────
    if ext.lstrip(".") == target_fmt:
        shutil.copy2(str(src), out_path)
        return out_path

    # ── Read source text ───────────────────────────────────────────────
    from adapters.document.libreoffice import LibreOfficeAdapter
    dummy = LibreOfficeAdapter()
    text = dummy.extract_text(str(src))

    # ── Write target ───────────────────────────────────────────────────
    creators: dict[str, Any] = {
        "odt": _create_odt,
        "docx": _create_docx,
        "pdf": _create_pdf,
        "txt": _create_txt,
        "html": _create_html,
    }
    if target_fmt in creators:
        creators[target_fmt](out_path, stem, text)
        return out_path

    raise ValueError(
        f"Unsupported target format {target_fmt!r}. "
        f"Supported: {sorted(creators)}"
    )


# ===================================================================
# Create helpers
# ===================================================================

def _create_txt(path: str, title: str, content: str) -> None:
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(f"{title}\n\n{content}")


def _create_html(path: str, title: str, content: str) -> None:
    escaped_content = _xml_escape(content)
    escaped_title = _xml_escape(title)
    html = (
        "<!DOCTYPE html>\n"
        "<html>\n<head>\n"
        f"<meta charset=\"utf-8\">\n<title>{escaped_title}</title>\n"
        "</head>\n<body>\n"
        f"<h1>{escaped_title}</h1>\n"
        f"<p>{escaped_content}</p>\n"
        "</body>\n</html>"
    )
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(html)


def _create_docx(path: str, title: str, content: str) -> None:
    """Create a .docx file using python-docx."""
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    # ── Default font ────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ── Title heading ───────────────────────────────────────────
    doc.add_heading(title, level=1)

    # ── Content paragraphs ──────────────────────────────────────
    for paragraph_text in content.split("\n"):
        paragraph_text = paragraph_text.strip()
        if not paragraph_text:
            doc.add_paragraph("")  # blank line
        else:
            doc.add_paragraph(paragraph_text)

    doc.save(path)


def _create_pdf(path: str, title: str, content: str) -> None:
    """Create a .pdf file using fpdf2."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("DejaVu", "", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
    pdf.add_font("DejaVu", "B", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
    pdf.set_font("DejaVu", "B", 16)
    pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("DejaVu", "", 11)
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            pdf.ln(4)
            continue
        # Ensure the line has characters that fit; use a safe encoding fallback
        safe = line.encode("latin-1", errors="replace").decode("latin-1")
        if not safe.strip():
            safe = " "
        # multi_cell with width w=0 (full page width) should always fit
        # If it still fails, skip the problematic line
        try:
            pdf.multi_cell(0, 6, safe)
        except Exception:
            pass  # silently skip unrenderable lines
    pdf.output(path)


def _create_odt(path: str, title: str, content: str) -> None:
    """Create a minimal valid ODF (ODT) file."""
    ns = (
        'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
        'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" '
        'xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" '
        'xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0"'
    )

    mimetype = b"application/vnd.oasis.opendocument.text"

    content_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        f'<office:document-content {ns} '
        'office:version="1.2">'
        '<office:body>'
        '<office:text>'
        f'<text:h text:style-name="Heading_20_1">{_xml_escape(title)}</text:h>'
        f'<text:p>{_xml_escape(content)}</text:p>'
        '</office:text>'
        '</office:body>'
        '</office:document-content>'
    )

    styles_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        f'<office:document-styles {ns} office:version="1.2">'
        '<office:styles>'
        '<style:style style:name="Heading_20_1" style:family="paragraph" '
        'style:parent-style-name="Standard" style:next-style-name="Text_20_body" '
        'style:default-outline-level="1">'
        '<style:text-properties fo:font-size="140%" fo:font-weight="bold"/>'
        '</style:style>'
        '</office:styles>'
        '</office:document-styles>'
    )

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", mimetype, compress_type=zipfile.ZIP_STORED)
        zf.writestr("content.xml", content_xml)
        zf.writestr("styles.xml", styles_xml)
        manifest = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<manifest:manifest '
            'xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" '
            'manifest:version="1.2">'
            '<manifest:file-entry '
            'manifest:media-type="application/vnd.oasis.opendocument.text" '
            'manifest:full-path="/"/>'
            '<manifest:file-entry manifest:media-type="text/xml" '
            'manifest:full-path="content.xml"/>'
            '<manifest:file-entry manifest:media-type="text/xml" '
            'manifest:full-path="styles.xml"/>'
            '</manifest:manifest>'
        )
        zf.writestr("META-INF/manifest.xml", manifest)


# ===================================================================
# Edit helpers
# ===================================================================

def _edit_docx(path: Path, content: str, insert_at: str | int | None) -> None:
    from docx import Document

    doc = Document(str(path))

    if insert_at is None and content:
        # Replace all body content
        for p in list(doc.paragraphs):
            p.clear()
        if doc.paragraphs:
            doc.paragraphs[0].text = content
        else:
            doc.add_paragraph(content)
    elif insert_at is not None and content:
        target = doc.paragraphs[-1] if insert_at in ("end", -1) else None
        try:
            idx = int(insert_at)  # type: ignore[arg-type]
            target = doc.paragraphs[idx] if idx < len(doc.paragraphs) else doc.paragraphs[-1]
        except (ValueError, IndexError):
            target = doc.paragraphs[-1]

        for line in content.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)

    doc.save(str(path))


def _edit_odt(path: Path, content: str, insert_at: str | int | None) -> None:
    """Edit an ODT by rewriting content.xml in the ZIP."""
    import zipfile

    content_new = content
    with zipfile.ZipFile(str(path), "r") as zf:
        old_content = zf.read("content.xml").decode("utf-8")

    if insert_at is None and content_new:
        # Simple regex replacement of text:p body
        # Replace the text inside the first text:p (below the heading)
        new_xml = re.sub(
            r"(<text:p>).*?(</text:p>)",
            lambda m: m.group(1) + _xml_escape(content_new) + m.group(2),
            old_content,
            count=1,
        )
    elif insert_at is not None and content_new:
        # Append new paragraph at the end before </office:text>
        new_para = f"<text:p>{_xml_escape(content_new)}</text:p>"
        new_xml = old_content.replace("</office:text>", new_para + "\n</office:text>")
    else:
        new_xml = old_content

    import zipfile as zfmod
    # Rewrite the zip in-place (careful: read all, then write)
    tmp = path.with_suffix(".odt.tmp")
    with zipfile.ZipFile(str(path), "r") as zin:
        with zipfile.ZipFile(str(tmp), "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "content.xml":
                    data = new_xml.encode("utf-8")
                zout.writestr(item, data)
    shutil.move(str(tmp), str(path))


def _edit_txt(path: Path, content: str, insert_at: str | int | None) -> None:
    if insert_at is None and content:
        path.write_text(content, encoding="utf-8")
    elif insert_at is not None and content:
        with open(str(path), "a", encoding="utf-8") as fh:
            fh.write(f"\n{content}\n")


# ===================================================================
# Extract helpers
# ===================================================================

def _extract_docx(src: Path) -> str:
    from docx import Document

    doc = Document(str(src))
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    return "\n".join(parts)


def _extract_odt(src: Path) -> str:
    """Extract text from an ODT by parsing content.xml."""
    import xml.etree.ElementTree as ET

    with zipfile.ZipFile(str(src), "r") as zf:
        content_xml = zf.read("content.xml")

    root = ET.fromstring(content_xml)
    ns = {
        "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
        "office": "urn:oasis:names:tc:opendocument:xmlns:office:1.0",
    }

    texts: list[str] = []
    for elem in root.iter():
        if elem.tag in (
            f"{{{ns['text']}}}p",
            f"{{{ns['text']}}}h",
        ):
            texts.append("".join(elem.itertext()))
    return "\n".join(texts)


def _extract_pdf(src: Path) -> str:
    """Extract text from a PDF using pdfminer.six (when available) or
    fall back to a basic binary heuristic."""
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        return pdfminer_extract(str(src))
    except ImportError:
        pass

    # Fallback: simple heuristic — extract text between parentheses
    text = src.read_bytes()
    results: list[str] = []
    in_text = False
    buf = ""
    for byte in text:
        ch = chr(byte)
        if ch == "(" and not in_text:
            in_text = True
            buf = ""
        elif ch == ")" and in_text:
            in_text = False
            if buf.strip():
                results.append(buf)
            buf = ""
        elif in_text:
            buf += ch
    return "\n".join(results)


def _extract_html(src: Path) -> str:
    import re

    text = src.read_text(encoding="utf-8", errors="replace")
    # Strip HTML tags
    text = re.sub(r"<[^>]+>", " ", text)
    # Collapse whitespace
    text = re.sub(r"\s+", " ", text).strip()
    return text


# ===================================================================
# Shared helpers
# ===================================================================

def _safe_filename(title: str) -> str:
    safe = "".join(c if c.isalnum() or c in " _.-\u2013" else "_" for c in title)
    return safe.strip() or "untitled"


def _xml_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )